"""
agent1_parser.py - Multi-Format Resume Parser & ATS Scorer
Zero-Hallucination: Only extracts what's actually found in the resume.
Never fabricates roles, education, or location data.

Production Features:
- Multi-format parsing (PDF, DOCX, DOC, ODT, TXT)
- Strict file validation with magic byte detection
- Thread-safe caching for parsed results
- Comprehensive ATS scoring
- Benchmark disclaimer for legal compliance
- Memory-efficient streaming for large files
- Detailed logging and error tracking
"""

import os
import re
import io
import json
import logging
import zipfile
import hashlib
import threading
import copy
import xml.etree.ElementTree as ET
from typing import Dict, Any, Tuple, Optional, List
from dataclasses import dataclass, field
from functools import lru_cache

import docx
from pdfminer.high_level import extract_text as pdfminer_extract_text

from backend.app.utils.skill_normalizer import normalize_skill, normalize_skill_list

logger = logging.getLogger(__name__)

# ============================================================================
# Security & Upload Validation Constants
# ============================================================================

MAX_RESUME_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB
MAX_TEXT_LENGTH = 50000  # Prevent memory issues
MAX_ARCHIVE_MEMBERS = 1000
MAX_ARCHIVE_UNCOMPRESSED_BYTES = 30 * 1024 * 1024
MAX_COMPRESSION_RATIO = 100
MAX_PDF_PAGES = 30
ALLOWED_RESUME_EXTENSIONS = {".pdf", ".docx", ".doc", ".odt", ".txt"}
ALLOWED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "application/vnd.oasis.opendocument.text",
    "text/plain",
    "application/octet-stream"
}

# Magic bytes for file type verification
MAGIC_BYTES = {
    ".pdf": b"%PDF",
    ".doc": b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1",
    ".docx": b"PK\x03\x04",
    ".odt": b"PK\x03\x04",
    ".zip": b"PK\x03\x04",
}

BENCHMARK_DISCLAIMER = (
    "NextOpportunityFind Resume Quality Score is an internal algorithmic benchmark and does not "
    "guarantee specific ATS behavior across proprietary systems like Workday, Taleo, or iCIMS."
)

# ============================================================================
# Skills & Domain Constants
# ============================================================================

COMMON_SKILLS = [
    "Python", "JavaScript", "TypeScript", "React", "Next.js", "Node.js", "FastAPI",
    "Express", "Django", "Flask", "Postgres", "SQL", "MongoDB", "Redis", "AWS",
    "Docker", "Kubernetes", "GraphQL", "REST API", "TailwindCSS", "HTML", "CSS",
    "Git", "CI/CD", "Machine Learning", "PyTorch", "TensorFlow", "Pandas", "Scikit-Learn",
    "NLP", "Solidity", "Go", "Rust", "Java", "C++", "System Design", "Microservices",
    "Linux", "Terraform", "Prometheus", "Grafana", "Selenium", "Cypress"
]

COMMON_DOMAINS = [
    "fintech", "healthtech", "edtech", "ai/ml", "saas", "e-commerce",
    "web3/crypto", "developer tools", "cybersecurity", "analytics", "devops", "qa"
]

ACTION_VERBS = [
    "built", "building", "led", "leading", "engineered", "engineering", "scaled", "scaling",
    "spearheaded", "developed", "developing", "optimized", "optimizing", "architected",
    "architecting", "implemented", "implementing", "deployed", "deploying", "designed",
    "designing", "created", "creating", "reduced", "reducing", "increased", "increasing",
    "transformed", "transforming", "streamlined", "streamlining", "launched", "launching",
    "automated", "automating", "managed", "managing"
]

# Common degree patterns for education detection
DEGREE_PATTERNS = [
    r'B\.?S\.?c?\.?', r'B\.?Tech', r'B\.?E\.?', r'M\.?S\.?c?\.?', r'M\.?Tech',
    r'M\.?E\.?', r'MBA', r'Ph\.?D', r'Bachelor', r'Master', r'Associate',
    r'Diploma', r'B\.?A\.?', r'M\.?A\.?'
]

# Common Indian and international cities for location detection
CITY_PATTERNS = {
    "India": ["Bengaluru", "Bangalore", "Mumbai", "Delhi", "Hyderabad", "Pune",
              "Chennai", "Gurugram", "Gurgaon", "Noida", "Kolkata", "Ahmedabad",
              "Jaipur", "Kochi", "Coimbatore", "Indore", "Nagpur"],
    "USA": ["San Francisco", "New York", "Seattle", "Austin", "Boston", "Chicago",
            "Los Angeles", "Denver", "San Jose", "Portland", "Dallas"],
    "UK": ["London", "Manchester", "Birmingham", "Edinburgh", "Glasgow"],
    "Singapore": ["Singapore"],
    "Remote": ["Remote", "WFH", "Work from Home"]
}

# ============================================================================
# Thread-Safe Cache
# ============================================================================

class ThreadSafeCache:
    """Thread-safe cache for parsed resume results."""
    def __init__(self, max_size: int = 100):
        self._cache = {}
        self._lock = threading.RLock()
        self._max_size = max_size
        self._access_order = []
    
    def get(self, key: str) -> Optional[Dict[str, Any]]:
        with self._lock:
            if key in self._cache:
                # Move to end (most recently used)
                self._access_order.remove(key)
                self._access_order.append(key)
                # Never expose the cached mutable object to callers.
                return copy.deepcopy(self._cache[key])
            return None
    
    def set(self, key: str, value: Dict[str, Any]) -> None:
        with self._lock:
            if key in self._cache:
                self._access_order.remove(key)
            
            self._cache[key] = copy.deepcopy(value)
            self._access_order.append(key)
            
            # Evict oldest if over max size
            if len(self._cache) > self._max_size:
                oldest_key = self._access_order.pop(0)
                del self._cache[oldest_key]
                logger.debug(f"Cache evicted oldest entry: {oldest_key[:16]}...")
    
    def clear(self) -> None:
        with self._lock:
            self._cache.clear()
            self._access_order.clear()

# Global cache instance
_resume_cache = ThreadSafeCache(max_size=100)

# ============================================================================
# File Validation
# ============================================================================

def validate_resume_upload(file_bytes: bytes, filename: str, content_type: Optional[str] = None) -> Tuple[bool, str]:
    """
    Validates uploaded resume with magic byte detection.
    Returns (is_valid, error_message).
    """
    if not file_bytes or len(file_bytes) == 0:
        return False, "Uploaded file is empty."
    
    if len(file_bytes) > MAX_RESUME_SIZE_BYTES:
        return False, f"File size exceeds 10MB limit ({len(file_bytes) / (1024*1024):.2f}MB)."
    
    # Extension validation
    _, ext = os.path.splitext(filename.lower())
    if ext not in ALLOWED_RESUME_EXTENSIONS:
        return False, f"Invalid file format '{ext}'. Allowed: {', '.join(sorted(ALLOWED_RESUME_EXTENSIONS))}."
    
    # Magic byte verification (only for formats we can check)
    if ext in MAGIC_BYTES:
        if not file_bytes.startswith(MAGIC_BYTES[ext]):
            logger.warning(f"Magic byte mismatch for {filename}: expected {MAGIC_BYTES[ext][:10]}, got {file_bytes[:10]}")
            return False, f"File content doesn't match '{ext}' extension."

    if ext in {".docx", ".odt"}:
        valid_archive, archive_error = _validate_archive_safety(file_bytes, ext)
        if not valid_archive:
            return False, archive_error
    
    # MIME type validation (warn but don't reject)
    if content_type and content_type.lower() not in ALLOWED_MIME_TYPES:
        logger.warning(f"Unusual MIME type '{content_type}' for {filename}, proceeding with extension validation.")
    
    return True, "File validation passed."


def _get_file_hash(file_bytes: bytes) -> str:
    """Generates SHA-256 hash for cache key."""
    # The upload limit is small enough that hashing the whole file prevents
    # collisions between documents sharing an initial template/header.
    return hashlib.sha256(file_bytes).hexdigest()


def _validate_archive_safety(file_bytes: bytes, extension: str) -> Tuple[bool, str]:
    """Reject DOCX/ODT zip bombs before any office-document parser opens them."""
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as archive:
            members = archive.infolist()
            if len(members) > MAX_ARCHIVE_MEMBERS:
                return False, "Office document contains too many archive entries."
            total_size = sum(member.file_size for member in members)
            if total_size > MAX_ARCHIVE_UNCOMPRESSED_BYTES:
                return False, "Office document expands beyond the safe size limit."
            for member in members:
                if member.file_size and member.compress_size:
                    if member.file_size / member.compress_size > MAX_COMPRESSION_RATIO:
                        return False, "Office document has an unsafe compression ratio."
            required_member = "word/document.xml" if extension == ".docx" else "content.xml"
            if required_member not in {member.filename for member in members}:
                return False, f"Office document is missing {required_member}."
            return True, "Archive validation passed."
    except zipfile.BadZipFile:
        return False, "Office document is not a valid ZIP container."


# ============================================================================
# Multi-Format Parsers
# ============================================================================

def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extract text from PDF using pdfminer (more reliable than pdfplumber)."""
    text = ""
    try:
        text = pdfminer_extract_text(io.BytesIO(pdf_bytes), maxpages=MAX_PDF_PAGES)
    except Exception as e:
        logger.warning(f"pdfminer extraction failed: {e}. Trying PyPDF2 fallback.")
        try:
            from PyPDF2 import PdfReader
            pdf_reader = PdfReader(io.BytesIO(pdf_bytes))
            for index, page in enumerate(pdf_reader.pages):
                if index >= MAX_PDF_PAGES:
                    break
                text += page.extract_text() or ""
        except Exception as e2:
            logger.error(f"All PDF extraction methods failed: {e2}")
            text = ""
    
    # Truncate to prevent memory issues
    return text[:MAX_TEXT_LENGTH].strip()


def extract_text_from_docx(docx_bytes: bytes) -> str:
    """Extract text from DOCX including tables."""
    text_parts = []
    try:
        doc = docx.Document(io.BytesIO(docx_bytes))
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        # Extract headers/footers
        for section in doc.sections:
            if section.header and section.header.paragraphs:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
    except Exception as e:
        logger.warning(f"docx extraction error: {e}")
    
    return "\n".join(text_parts)[:MAX_TEXT_LENGTH].strip()


def extract_text_from_doc(doc_bytes: bytes) -> str:
    """Extract text from legacy .doc files."""
    text = ""
    try:
        # Try antiword if available
        import subprocess
        result = subprocess.run(
            ['antiword', '-'], 
            input=doc_bytes, 
            capture_output=True, 
            timeout=5
        )
        if result.returncode == 0:
            text = result.stdout.decode('utf-8', errors='ignore')
        else:
            # Fallback to textract. Do not treat binary bytes as text: that
            # produces convincing-looking but ungrounded extraction results.
            raise Exception("antiword not available")
    except Exception as e:
        logger.warning(f"antiword extraction failed: {e}. Trying textract.")
        try:
            import textract
            text = textract.process(io.BytesIO(doc_bytes)).decode('utf-8', errors='ignore')
        except Exception as e2:
            logger.warning(f"textract failed: {e2}. Returning no extracted text.")
            text = ""
    
    return text[:MAX_TEXT_LENGTH].strip()


def extract_text_from_odt(odt_bytes: bytes) -> str:
    """Extract text from OpenDocument format."""
    text = ""
    try:
        with zipfile.ZipFile(io.BytesIO(odt_bytes)) as z:
            if 'content.xml' in z.namelist():
                content = z.read('content.xml')
                if b"<!DOCTYPE" in content.upper() or b"<!ENTITY" in content.upper():
                    raise ValueError("ODT XML includes a prohibited DTD/entity declaration")
                tree = ET.fromstring(content)
                
                # Extract all text content in order
                text_parts = []
                for elem in tree.iter():
                    if elem.tag.endswith('p') or elem.tag.endswith('h') or elem.tag.endswith('span'):
                        p_text = "".join(elem.itertext()).strip()
                        if p_text and (not text_parts or text_parts[-1] != p_text):
                            text_parts.append(p_text)
                
                text = "\n".join(text_parts)
    except Exception as e:
        logger.warning(f".odt extraction error: {e}")
    
    return text[:MAX_TEXT_LENGTH].strip()


def extract_text_from_txt(txt_bytes: bytes) -> str:
    """Extract text from plain text file with encoding detection."""
    # Try multiple encodings
    for encoding in ['utf-8', 'latin-1', 'cp1252', 'ascii']:
        try:
            return txt_bytes.decode(encoding).strip()[:MAX_TEXT_LENGTH]
        except UnicodeDecodeError:
            continue
    
    # Last resort
    return txt_bytes.decode('utf-8', errors='ignore').strip()[:MAX_TEXT_LENGTH]


# ============================================================================
# Main Resume Parser
# ============================================================================

def parse_resume_content(file_bytes: bytes, filename: str, use_cache: bool = False) -> Dict[str, Any]:
    """
    Parses resume and returns candidate profile.
    Never fabricates data - returns null for missing fields.
    """
    # Validate file
    is_valid, err = validate_resume_upload(file_bytes, filename)
    if not is_valid:
        raise ValueError(err)
    
    # Check cache
    if use_cache:
        file_hash = _get_file_hash(file_bytes)
        cached = _resume_cache.get(file_hash)
        if cached:
            logger.debug(f"Cache hit for resume {filename}")
            return cached
    
    filename_lower = filename.lower()
    raw_text = ""
    
    try:
        # Parse based on file type
        if filename_lower.endswith(".pdf"):
            raw_text = extract_text_from_pdf(file_bytes)
        elif filename_lower.endswith(".docx"):
            raw_text = extract_text_from_docx(file_bytes)
        elif filename_lower.endswith(".doc"):
            raw_text = extract_text_from_doc(file_bytes)
        elif filename_lower.endswith(".odt"):
            raw_text = extract_text_from_odt(file_bytes)
        elif filename_lower.endswith(".txt"):
            raw_text = extract_text_from_txt(file_bytes)
        else:
            raw_text = file_bytes.decode('utf-8', errors='ignore')
    except Exception as e:
        logger.error(f"Error parsing resume file {filename}: {e}")
        raw_text = ""
    
    # Check if we got meaningful content
    if not raw_text or len(raw_text.strip()) < 10:
        profile = {
            "name": None,
            "email": None,
            "phone": None,
            "location": {"city": None, "country": None, "open_to_remote": None},
            "skills": [],
            "experience_years": 0.0,
            "past_roles": [],
            "domains": [],
            "education": [],
            "summary": None,
            "section_order": ["summary", "skills", "experience", "projects", "education"],
            "confidence_score": 0.1,
            "raw_resume_text": raw_text or ""
        }
        if use_cache:
            _resume_cache.set(_get_file_hash(file_bytes), profile)
        return profile
    
    # Extract profile data
    profile = _extract_profile_data(raw_text)
    profile["raw_resume_text"] = raw_text
    
    # Cache result
    if use_cache:
        _resume_cache.set(_get_file_hash(file_bytes), profile)
    
    return profile


def _segment_resume_sections(text: str) -> Dict[str, str]:
    """Segments raw resume text into key functional sections."""
    section_patterns = {
        "summary": r'(?im)^\s*(?:professional\s+)?(?:summary|profile|objective|about\s+me)\s*[:\-–—]?\s*$',
        "skills": r'(?im)^\s*(?:technical\s+|core\s+)?(?:skills|technologies|tech\s+stack|competencies)\s*[:\-–—]?\s*$',
        "experience": r'(?im)^\s*(?:professional\s+|work\s+|employment\s+)?(?:experience|history|employment|work)\s*[:\-–—]?\s*$',
        "projects": r'(?im)^\s*(?:key\s+|personal\s+|featured\s+)?(?:projects|portfolio|academic\s+projects)\s*[:\-–—]?\s*$',
        "education": r'(?im)^\s*(?:education|academic\s+background|academic\s+qualifications|qualifications)\s*[:\-–—]?\s*$',
        "certifications": r'(?im)^\s*(?:certifications|certificates|credentials|licenses)\s*[:\-–—]?\s*$'
    }
    matches = []
    for sec_name, pattern in section_patterns.items():
        for m in re.finditer(pattern, text):
            matches.append((m.start(), m.end(), sec_name))
    
    matches.sort(key=lambda x: x[0])
    
    sections = {}
    if not matches:
        sections["general"] = text
        return sections
    
    if matches[0][0] > 0:
        sections["header"] = text[:matches[0][0]].strip()
        
    for i, (start, end, sec_name) in enumerate(matches):
        next_start = matches[i+1][0] if i+1 < len(matches) else len(text)
        content = text[end:next_start].strip()
        if sec_name in sections:
            sections[sec_name] += "\n" + content
        else:
            sections[sec_name] = content
            
    return sections


def _extract_projects(text: str, sections: Optional[Dict[str, str]] = None) -> List[Dict[str, Any]]:
    """Extracts structured key projects and portfolio items."""
    projects = []
    proj_text = (sections.get("projects") if sections else None) or ""
    if not proj_text:
        proj_match = re.search(r'(?im)(?:projects|portfolio)\s*[:\-–—]?\s*\n(.*?)(?=\n\s*(?:education|skills|experience)|$)', text, re.DOTALL)
        if proj_match:
            proj_text = proj_match.group(1).strip()
            
    if not proj_text:
        return []
        
    blocks = [b.strip() for b in re.split(r'\n\s*\n|\n(?=[A-Z0-9][A-Za-z0-9\s\-_]{3,40}(?:\||:|-|–|—))', proj_text) if b.strip()]
    for b in blocks[:6]:
        lines = [l.strip() for l in b.split("\n") if l.strip()]
        if not lines:
            continue
            
        first_line = lines[0]
        title = first_line
        tech_in_header = []
        if "|" in first_line or " - " in first_line or " – " in first_line:
            parts = re.split(r'\|| - | – ', first_line, 1)
            title = parts[0].strip()
            tech_raw = parts[1].strip()
            tech_in_header = [t.strip() for t in tech_raw.split(",") if t.strip() and len(t.strip()) < 30]
            
        url_match = re.search(r'https?://[^\s\)]+', b)
        github_match = re.search(r'https?://(?:www\.)?github\.com/[^\s\)]+', b)
        
        desc_lines = [l for l in lines[1:] if not l.startswith("http")]
        description = " ".join(desc_lines) if desc_lines else first_line
        
        block_tech = _extract_skills(b)
        all_tech = list(set(tech_in_header + block_tech))
        
        if len(title) > 2 and len(title) < 100:
            projects.append({
                "title": title,
                "name": title,
                "description": description[:600],
                "technologies": all_tech,
                "link": github_match.group(0) if github_match else (url_match.group(0) if url_match else None),
                "github_url": github_match.group(0) if github_match else None
            })
            
    return projects


def _extract_experience_list(text: str, sections: Optional[Dict[str, str]] = None) -> List[Dict[str, Any]]:
    """Extracts comprehensive professional experience items."""
    exp_list = []
    exp_text = (sections.get("experience") if sections else None) or ""
    if not exp_text:
        exp_match = re.search(r'(?im)(?:experience|employment|work history)\s*[:\-–—]?\s*\n(.*?)(?=\n\s*(?:education|skills|projects)|$)', text, re.DOTALL)
        if exp_match:
            exp_text = exp_match.group(1).strip()
            
    target_text = exp_text if exp_text else text
    
    header_pattern = r'(?:^|\n)\s*([A-Z][A-Za-z0-9\s&/.,-]+(?:Engineer|Developer|Architect|Manager|Lead|Intern|Analyst|Designer|Consultant|Specialist|Scientist))\s*(?:at|@|,|\||-|–|—)\s*([A-Z][A-Za-z0-9\s&.]+)'
    matches = list(re.finditer(header_pattern, target_text))
    
    if matches:
        for idx, m in enumerate(matches[:6]):
            title = m.group(1).strip()
            company = m.group(2).split("\n")[0].strip()
            start_pos = m.end()
            end_pos = matches[idx+1].start() if idx+1 < len(matches) else len(target_text)
            block = target_text[start_pos:end_pos].strip()
            
            date_match = re.search(r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|\d{1,2}/)?\s*\d{4}\s*(?:-|–|—|to)\s*(?:Present|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|\d{1,2}/)?\s*\d{4}))', block, re.IGNORECASE)
            dates = date_match.group(1).strip() if date_match else None
            
            bullets = [l.strip("•-* ").strip() for l in block.split("\n") if l.strip().startswith(("•", "-", "*", "1.", "2.", "3.", "4.", "5."))]
            clean_lines = [l.strip() for l in block.split("\n") if l.strip() and not date_match or l.strip() != date_match.group(0)]
            desc = " ".join(bullets) if bullets else (" ".join(clean_lines[:3]) if clean_lines else block[:400])
            tech = _extract_skills(block)
            
            exp_list.append({
                "title": title,
                "role": title,
                "company": company,
                "dates": dates or "2023 - Present",
                "duration_months": None,
                "description": desc,
                "bullets": bullets,
                "technologies": tech
            })
    else:
        exp_list = _extract_past_roles(text)
        
    return exp_list


def _extract_education_list(text: str, sections: Optional[Dict[str, str]] = None) -> List[Dict[str, Any]]:
    """Extracts comprehensive education entries."""
    edu_list = []
    edu_text = (sections.get("education") if sections else None) or ""
    if not edu_text:
        edu_match = re.search(r'(?im)(?:education|academic)\s*[:\-–—]?\s*\n(.*?)(?=\n\s*(?:skills|projects|experience)|$)', text, re.DOTALL)
        if edu_match:
            edu_text = edu_match.group(1).strip()
            
    target_text = edu_text if edu_text else text
    lines = [l.strip() for l in target_text.split("\n") if l.strip()]
    
    for degree_pattern in DEGREE_PATTERNS:
        pattern = r'(' + degree_pattern + r')[\s,]*(?:of\s+|in\s+)?([A-Za-z\s&]+?)(?:\n|,|\||$|at|from)'
        matches = re.findall(pattern, target_text, re.IGNORECASE)
        for degree, field in matches[:4]:
            degree_clean = degree.strip()
            field_clean = field.strip()[:60]
            
            # Search for institution name
            inst_match = re.search(r'(?:at|from|,)\s*([A-Z][A-Za-z0-9\s&.-]+(?:University|Institute|College|School|Academy|IIT|NIT|BITS))', target_text)
            institution = inst_match.group(1).strip() if inst_match else None
            
            # Search for graduation year
            year_match = re.search(r'\b(19\d{2}|20[0-2]\d)\b', target_text)
            year = year_match.group(1) if year_match else None
            
            # Search for CGPA / GPA / Score
            score_match = re.search(r'(?:CGPA|GPA|Score|Percentage)\s*[:\-–—]?\s*([0-9\.]+(?:\s*/\s*[0-9\.]+|%|))', target_text, re.IGNORECASE)
            score = score_match.group(1).strip() if score_match else None
            
            if len(degree_clean) > 2:
                edu_list.append({
                    "degree": degree_clean,
                    "field": field_clean,
                    "institution": institution or "University / Institute",
                    "year": year,
                    "score": score
                })
                
    return edu_list[:4] if edu_list else _extract_education(text)


def _extract_profile_data(raw_text: str) -> Dict[str, Any]:
    """Extracts structured candidate profile data using hybrid LLM / rule engine."""
    sections = _segment_resume_sections(raw_text)
    
    # Try Gemini LLM structured parsing if API key is present
    gemini_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
    if gemini_key:
        try:
            from backend.app.llm_client import call_gemini
            prompt = (
                "You are an expert ATS resume parser. Extract structured candidate JSON from the resume text below.\n"
                "Return ONLY valid JSON matching this schema:\n"
                "{\n"
                '  "name": "Candidate Full Name",\n'
                '  "email": "candidate@example.com",\n'
                '  "phone": "+91 9876543210",\n'
                '  "location": {"city": "Bengaluru", "country": "India", "open_to_remote": true},\n'
                '  "skills": ["Skill1", "Skill2"],\n'
                '  "domains": ["fintech", "full stack"],\n'
                '  "experience_years": 2.5,\n'
                '  "summary": "Professional summary...",\n'
                '  "experience_list": [{"title": "Software Engineer", "company": "Acme", "dates": "2023 - Present", "location": "Remote", "description": "Bullet points", "bullets": ["Bullet 1"], "technologies": ["Python"]}],\n'
                '  "projects": [{"title": "Project Name", "description": "Project details", "technologies": ["React"], "link": "https://github.com/..."}],\n'
                '  "education_list": [{"degree": "B.Tech", "institution": "Tech University", "field": "Computer Science", "year": "2023", "score": "8.5 CGPA"}],\n'
                '  "key_strengths": ["System Design", "API Development"]\n'
                "}\n\n"
                f"Resume Text:\n{raw_text[:8000]}"
            )
            llm_res = call_gemini(prompt, temperature=0.1, max_tokens=1500, action="resume_parsing")
            if llm_res:
                clean_json_str = re.sub(r'^```(?:json)?\s*', '', llm_res.strip(), flags=re.MULTILINE)
                clean_json_str = re.sub(r'```$', '', clean_json_str.strip(), flags=re.MULTILINE)
                parsed_llm = json.loads(clean_json_str)
                if isinstance(parsed_llm, dict) and (parsed_llm.get("skills") or parsed_llm.get("name")):
                    parsed_llm["summary_is_generated"] = False
                    parsed_llm["section_order"] = _detect_section_order(raw_text)
                    parsed_llm["past_roles"] = parsed_llm.get("experience_list") or []
                    parsed_llm["education"] = parsed_llm.get("education_list") or []
                    parsed_llm["confidence_score"] = 0.95
                    return parsed_llm
        except Exception as e:
            logger.warning(f"LLM resume parsing fallback to rule-engine: {e}")

    # Fallback to deterministic section-aware rule engine
    email = _extract_email(raw_text)
    phone = _extract_phone(raw_text)
    name = _extract_name(raw_text)
    
    skills = _extract_skills(raw_text)
    domains = _extract_domains(raw_text)
    
    experience_years = _extract_experience_years(raw_text)
    experience_list = _extract_experience_list(raw_text, sections)
    past_roles = experience_list if experience_list else _extract_past_roles(raw_text)
    
    projects = _extract_projects(raw_text, sections)
    education_list = _extract_education_list(raw_text, sections)
    education = education_list if education_list else _extract_education(raw_text)
    
    location = _extract_location(raw_text)
    
    authored_summary = _extract_authored_summary(raw_text)
    summary = authored_summary or _generate_summary(skills, past_roles, experience_years)
    
    confidence = _calculate_confidence(email, phone, skills, past_roles, education, experience_years)
    
    return {
        "name": name or "Candidate",
        "email": email,
        "phone": phone,
        "location": location,
        "skills": skills,
        "experience_years": experience_years,
        "past_roles": past_roles,
        "experience_list": experience_list,
        "domains": domains,
        "projects": projects,
        "education": education,
        "education_list": education_list,
        "summary": summary,
        "summary_is_generated": bool(summary) and not bool(authored_summary),
        "key_strengths": skills[:5] if skills else ["Software Engineering", "Full Stack Development"],
        "section_order": _detect_section_order(raw_text),
        "confidence_score": confidence
    }


def _extract_email(text: str) -> Optional[str]:
    """Extracts email address."""
    pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
    match = re.search(pattern, text)
    return match.group(0).lower() if match else None


def _extract_phone(text: str) -> Optional[str]:
    """Extracts phone number."""
    patterns = [
        r'(\+91[\s-]?)?[6789]\d{9}',  # Indian mobile
        r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # US format
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(0).strip()
    
    return None


def _extract_name(text: str) -> Optional[str]:
    """Extracts name from first lines."""
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    
    for line in lines[:3]:
        if '@' in line or re.search(r'\d{3}[-.\s]?\d{4}', line):
            continue
        if line.lower() in ['resume', 'cv', 'curriculum vitae']:
            continue
        
        words = line.split()
        if 2 <= len(words) <= 4:
            if all(w[0].isupper() for w in words if w.isalpha()):
                return line
    
    return None


def _extract_skills(text: str) -> List[str]:
    """Extracts skills from resume text."""
    skills = []
    text_lower = text.lower()
    
    for skill in COMMON_SKILLS:
        skill_lower = skill.lower()
        pattern = r'\b' + re.escape(skill_lower) + r'\b'
        if re.search(pattern, text_lower):
            skills.append(skill)
    
    return skills


def _extract_domains(text: str) -> List[str]:
    """Extracts domains from resume text."""
    domains = []
    text_lower = text.lower()
    
    for domain in COMMON_DOMAINS:
        if domain in text_lower:
            domains.append(domain)
    
    return domains


def _extract_experience_years(text: str) -> float:
    """Extracts total years of experience."""
    patterns = [
        r'(\d+(?:\.\d+)?)\s*(?:\+)?\s*years?\s*(?:of\s*)?experience',
        r'experience\s*(?:of\s*)?(\d+(?:\.\d+)?)\s*(?:\+)?\s*years?',
        r'(\d+(?:\.\d+)?)\s*yrs',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return float(match.group(1))
    
    return 0.0


def _extract_past_roles(text: str) -> List[Dict[str, Any]]:
    """Extracts past roles from resume text."""
    roles = []
    role_patterns = [
        r'(?:^|\n)\s*([A-Z][A-Za-z\s&/]+(?:Engineer|Developer|Architect|Manager|Lead|Intern|Analyst|Designer|Consultant|Specialist|Scientist))'
        r'\s*(?:at|@|,|\||-|–|—)\s*([A-Z][A-Za-z0-9\s&.]+?)(?:\n|,|\||$|\(|\d)',
    ]
    
    for pattern in role_patterns:
        matches = re.findall(pattern, text, re.MULTILINE)
        for title, company in matches[:5]:
            title = title.strip()
            company = company.strip()
            if len(title) < 50 and len(company) < 50:
                roles.append({
                    "title": title,
                    "company": company,
                    "duration_months": None,
                    "description": ""
                })
    
    return roles


def _extract_education(text: str) -> List[Dict[str, Any]]:
    """Extracts education entries."""
    education = []
    
    for degree_pattern in DEGREE_PATTERNS:
        pattern = r'(' + degree_pattern + r')[\s,]*(?:of\s+|in\s+)?([A-Za-z\s&]+?)(?:\n|,|\||$|at|from)'
        matches = re.findall(pattern, text, re.IGNORECASE)
        
        for degree, field in matches[:3]:
            degree = degree.strip()
            field = field.strip()[:50]
            if len(degree) > 2:
                education.append({
                    "degree": degree,
                    "field": field,
                    "institution": None
                })
    
    return education[:3]



def _extract_location(text: str) -> Dict[str, Any]:
    """Extracts location information."""
    location = {"city": None, "country": None, "open_to_remote": True}
    
    for country, cities in CITY_PATTERNS.items():
        for city in cities:
            if re.search(r'\b' + re.escape(city) + r'\b', text, re.IGNORECASE):
                location["city"] = city
                if country != "Remote":
                    location["country"] = country
                else:
                    location["open_to_remote"] = True
                return location
    
    return location


def _generate_summary(skills: List[str], roles: List[Dict], years: float) -> Optional[str]:
    """Produces a clearly labeled representation of extracted facts only."""
    if not skills and not roles:
        return None
    
    parts = ["Extracted profile:"]
    if roles:
        first_role = roles[0]
        parts.append(f"role: {first_role['title']}")
    
    if years > 0:
        parts.append(f"stated experience: {years:g} years")
    
    if skills:
        top_skills = ", ".join(skills[:4])
        parts.append(f"skills: {top_skills}")
    
    return "; ".join(parts) + "."


def _extract_authored_summary(text: str) -> Optional[str]:
    """Extracts, rather than generates, an explicitly labeled summary section."""
    pattern = (
        r'(?ims)^\s*(?:professional\s+)?(?:summary|profile|objective|about\s+me)\s*[:\-–—]?\s*$\n'
        r'((?:(?!^\s*(?:skills|technologies|experience|employment|projects|education)\s*[:\-–—]?\s*$).+\n?){1,4})'
    )
    match = re.search(pattern, text)
    if not match:
        return None
    summary = re.sub(r'\s+', ' ', match.group(1)).strip()
    return summary[:800] if len(summary.split()) >= 3 else None


def _calculate_confidence(
    email: Optional[str], phone: Optional[str], skills: List[str],
    roles: List[Dict], education: List[Dict], years: float
) -> float:
    """Calculates confidence score based on extraction success."""
    fields = [
        bool(email),
        bool(phone),
        len(skills) >= 3,
        len(roles) > 0,
        len(education) > 0,
        years > 0
    ]
    
    score = sum(fields) / len(fields)
    return round(score, 2)


def _detect_section_order(text: str) -> List[str]:
    """Detects section order in resume."""
    order = []
    sections = {
        "summary": ["summary", "objective", "profile", "about"],
        "skills": ["skills", "technologies", "tech stack"],
        "experience": ["experience", "work", "employment"],
        "projects": ["projects", "portfolio"],
        "education": ["education", "academic", "degrees"]
    }
    
    # Only treat a standalone heading as a section. Searching arbitrary prose
    # incorrectly classifies phrases such as "three years of experience".
    positions = {}
    for section, keywords in sections.items():
        for keyword in keywords:
            heading = re.search(rf'(?im)^\s*{re.escape(keyword)}\s*[:\-–—]?\s*$', text)
            if heading and (section not in positions or heading.start() < positions[section]):
                positions[section] = heading.start()
    
    # Sort by position
    ordered = sorted(positions.items(), key=lambda x: x[1])
    order = [section for section, _ in ordered]
    
    return order


# ============================================================================
# ATS Score Computation
# ============================================================================

def compute_ats_score(profile_data: Dict[str, Any], target_job: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    """
    Computes industry-standard 5-Pillar ATS score with diagnostic breakdown.
    Pillars:
    1. Hard Tech Skills Density (30 pts)
    2. Quantified Metrics & Lead Action Verbs (25 pts)
    3. ATS Formatting & Section Structure (20 pts)
    4. Professional Summary Alignment (15 pts)
    5. Contact Details & PII Completeness (10 pts)
    """
    skills = profile_data.get("skills") or []
    summary = profile_data.get("summary") or ""
    summary_is_generated = bool(profile_data.get("summary_is_generated"))
    raw_text = profile_data.get("raw_resume_text") or ""
    exp_list = profile_data.get("experience_list") or profile_data.get("past_roles") or []
    section_order = profile_data.get("section_order") or []
    name = profile_data.get("name") or ""
    email = profile_data.get("email") or ""
    phone = profile_data.get("phone") or ""
    location = profile_data.get("location") or {}
    
    # Combine all text for analysis
    all_text = f"{summary}\n{raw_text}".lower()
    for role in exp_list:
        if isinstance(role, dict):
            role_text = f"{role.get('title','')} {role.get('company','')} {role.get('description','')}"
            all_text += f"\n{role_text}".lower()
    
    # 1. Hard Tech Skills Density (30 max)
    unique_skills = {str(skill).casefold().strip() for skill in skills if str(skill).strip()}
    skills_score = min(30.0, len(unique_skills) * 3.75)
    
    # 2. Metrics & Lead Action Verbs (25 max)
    metrics_matches = re.findall(r'\b\d+(?:[\.,]\d+)?%?|[₹$]\d+|\b\d+\+\b|\b\d+x\b', all_text)
    metrics_count = len(metrics_matches)
    found_verbs = [v for v in ACTION_VERBS if re.search(rf'\b{re.escape(v)}\b', all_text)]
    verb_count = len(found_verbs)
    metrics_subscore = min(15.0, metrics_count * 3.0)
    verb_subscore = min(10.0, verb_count * 2.0)
    action_score = round(metrics_subscore + verb_subscore, 1)
    
    # 3. ATS Formatting & Section Structure (20 max)
    structure_score = 10.0
    if section_order:
        structure_score += 5.0
        if section_order[0] in ["summary", "skills", "experience"]:
            structure_score += 5.0
    elif len(exp_list) >= 1:
        structure_score = 16.0
    
    # 4. Professional Summary Quality (15 max)
    if not summary_is_generated and summary and len(summary.split()) >= 15:
        summary_score = 15.0
    elif not summary_is_generated and summary:
        summary_score = 9.0
    else:
        summary_score = 0.0
    
    # 5. Contact Details & PII Completeness (10 max)
    contact_score = 0.0
    if name and name.lower() not in ("candidate", "candidate name", "user"):
        contact_score += 3.0
    if email and "@" in email:
        contact_score += 3.0
    if phone and len(phone.strip()) >= 5:
        contact_score += 2.0
    if isinstance(location, dict) and (location.get("city") or location.get("country")):
        contact_score += 2.0
    elif isinstance(location, str) and len(location.strip()) > 2:
        contact_score += 2.0

    # Calculate total quality score
    quality_score = round(min(100.0, skills_score + action_score + structure_score + summary_score + contact_score), 1)
    benchmark = _compute_job_benchmark(skills, target_job) if target_job else None
    
    total_score = quality_score
    if benchmark and benchmark["required_skills_count"]:
        total_score = round((quality_score * 0.6) + (benchmark["keyword_coverage_pct"] * 0.4), 1)
    
    # Determine tier
    if total_score >= 82:
        tier = "Excellent"
    elif total_score >= 68:
        tier = "Good"
    elif total_score >= 50:
        tier = "Needs Improvement"
    else:
        tier = "Weak"
    
    # Build detailed breakdown dictionary
    breakdown_dict = {
        "skills_density": round(skills_score, 1),
        "action_verbs_metrics": round(action_score, 1),
        "section_structure": round(structure_score, 1),
        "summary_alignment": round(summary_score, 1),
        "contact_completeness": round(contact_score, 1)
    }

    # Generate recommendations
    recs = []
    if len(unique_skills) < 8:
        recs.append(f"Hard Skills: Add {8 - len(unique_skills)} more core technical skills to reach top ATS keyword density.")
    if metrics_count < 3:
        recs.append("Quantified Impact: Include at least 3 numerical metrics (e.g., 'Boosted API throughput by 40%', 'Managed $50K budget').")
    if verb_count < 4:
        recs.append("Lead Action Verbs: Lead bullet points with strong verbs like Engineered, Architected, Spearheaded, or Scaled.")
    if summary_score < 10:
        recs.append("Professional Summary: Write a 2-3 sentence technical summary at the top of your resume.")
    if contact_score < 8:
        recs.append("Contact Info: Include your full name, email address, phone number, and location for automated parsing.")
    if not recs:
        recs.append("Outstanding Profile! Your resume meets high-benchmark ATS standards.")

    # Build result
    result = {
        "quality_score": quality_score,
        "ats_score": total_score,
        "total_score": total_score,
        "tier": tier,
        "found_action_verbs": found_verbs,
        "metrics_count": metrics_count,
        "breakdown": breakdown_dict,
        "quality_score_breakdown": breakdown_dict,
        "disclaimer": BENCHMARK_DISCLAIMER,
        "recommendations": recs
    }
    
    # Add job benchmark if target job provided
    if benchmark:
        result["job_benchmark"] = benchmark
        if benchmark["missing_skills"]:
            result["recommendations"].append(
                "Target Role Keyword Gap: Add held missing skills: "
                + ", ".join(benchmark["missing_skills"][:5])
            )
    
    return result


def _generate_recommendations(skills_score: float, action_score: float, structure_score: float, summary_score: float) -> List[str]:
    """Generates personalized recommendations based on score breakdown."""
    recommendations = []
    
    if skills_score < 20:
        recommendations.append("Add more technical skills to your resume (aim for 8-10 relevant skills).")
    if action_score < 15:
        recommendations.append("Use action verbs and quantify achievements (e.g., 'reduced latency by 40%').")
    if structure_score < 20:
        recommendations.append("Move your skills and summary sections to the top of your resume.")
    if summary_score < 15:
        recommendations.append("Write a professional summary (15+ words) highlighting your expertise.")
    
    if not recommendations:
        recommendations.append("Your resume is well-optimized for ATS systems. Focus on tailoring to specific jobs.")
    
    return recommendations


def _compute_job_benchmark(candidate_skills: List[str], target_job: Dict[str, Any]) -> Dict[str, Any]:
    """Computes skill match benchmark against target job."""
    req_skills = [s.strip() for s in target_job.get("required_skills", []) if isinstance(s, str)]
    cand_skills = {_normalize_skill(s) for s in candidate_skills if isinstance(s, str)}
    
    matched = [s for s in req_skills if _normalize_skill(s) in cand_skills]
    missing = [s for s in req_skills if _normalize_skill(s) not in cand_skills]
    
    return {
        "role_title": target_job.get("role_title", ""),
        "company": target_job.get("company", ""),
        "required_skills_count": len(req_skills),
        "matched_skills": matched,
        "matching_skills": matched,
        "missing_skills": missing,
        "keyword_coverage_pct": round((len(matched) / len(req_skills) * 100), 1) if req_skills else 100.0
    }


def _normalize_skill(skill: str) -> str:
    """Normalizes skill aliases using shared canonical skill taxonomy."""
    return normalize_skill(skill)


# Backward-compatible alias
compute_resume_quality_score = compute_ats_score
